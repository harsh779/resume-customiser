"""
Extract plain text from a resume file (DOCX or PDF).
Usage: python parse_resume.py <file_path>
Outputs extracted text to stdout.
"""
import sys
import os


def parse_docx(path):
    from docx import Document
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return "\n".join(lines)


def parse_pdf(path):
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n".join(pages)
    except ImportError:
        import PyPDF2
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n".join(pages)


def main():
    if len(sys.argv) < 2:
        print("Usage: parse_resume.py <file_path>", file=sys.stderr)
        sys.exit(1)

    path = sys.argv[1]

    if not os.path.exists(path):
        print(f"Error: File not found: {path}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        text = parse_docx(path)
    elif ext == ".pdf":
        text = parse_pdf(path)
    elif ext == ".txt":
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        print(f"Error: Unsupported file type '{ext}'. Use DOCX, PDF, or TXT.", file=sys.stderr)
        sys.exit(1)

    if not text.strip():
        print("Error: No text extracted from resume. File may be image-based or empty.", file=sys.stderr)
        sys.exit(1)

    print(text)


if __name__ == "__main__":
    main()
