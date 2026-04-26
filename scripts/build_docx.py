"""
Build an ATS-safe single-column DOCX resume from structured JSON.
Usage: python build_docx.py <json_file_path> <output_docx_path>

JSON schema:
{
  "name": "Full Name",
  "contact": {
    "email": "...", "phone": "...", "linkedin": "...",
    "location": "...", "github": "..."
  },
  "summary": "...",
  "experience": [
    {
      "title": "...", "company": "...", "location": "...",
      "dates": "...", "bullets": ["...", "..."]
    }
  ],
  "skills": {
    "Technical": ["..."], "Tools": ["..."], "Domain": ["..."]
  },
  "education": [
    {"degree": "...", "institution": "...", "dates": "...", "notes": "..."}
  ],
  "certifications": ["..."],
  "gaps": ["..."]
}
"""
import sys
import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_section_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text.upper())
    set_font(run, size=11, bold=True, color=(31, 73, 125))

    # Horizontal rule under heading
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(text)
    set_font(run, size=10.5)
    return para


def add_role_header(doc, title, company, location, dates):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(1)

    title_run = para.add_run(title)
    set_font(title_run, size=11, bold=True)

    if company:
        sep_run = para.add_run("  |  ")
        set_font(sep_run, size=11)
        co_run = para.add_run(company)
        set_font(co_run, size=11)

    if location:
        loc_run = para.add_run(f"  ·  {location}")
        set_font(loc_run, size=10, color=(89, 89, 89))

    if dates:
        date_run = para.add_run(f"  |  {dates}")
        set_font(date_run, size=10, color=(89, 89, 89))

    return para


def set_page_margins(doc, top=0.75, bottom=0.75, left=0.9, right=0.9):
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def build(data, output_path):
    doc = Document()
    set_page_margins(doc)

    # Remove default styles that add unwanted spacing
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- NAME ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(data.get("name", ""))
    set_font(name_run, size=18, bold=True)

    # --- CONTACT ---
    contact = data.get("contact", {})
    contact_parts = []
    for field in ["location", "phone", "email", "linkedin", "github"]:
        val = contact.get(field, "").strip()
        if val:
            contact_parts.append(val)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(6)
        contact_run = contact_para.add_run("  |  ".join(contact_parts))
        set_font(contact_run, size=10, color=(89, 89, 89))

    # --- SUMMARY ---
    summary = data.get("summary", "").strip()
    if summary:
        add_section_heading(doc, "Professional Summary")
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(summary)
        set_font(run, size=10.5)

    # --- EXPERIENCE ---
    experience = data.get("experience", [])
    if experience:
        add_section_heading(doc, "Experience")
        for role in experience:
            add_role_header(
                doc,
                role.get("title", ""),
                role.get("company", ""),
                role.get("location", ""),
                role.get("dates", ""),
            )
            for bullet in role.get("bullets", []):
                add_bullet(doc, bullet)

    # --- SKILLS ---
    skills = data.get("skills", {})
    if skills:
        add_section_heading(doc, "Skills")
        for category, items in skills.items():
            if items:
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(1)
                label_run = para.add_run(f"{category}: ")
                set_font(label_run, size=10.5, bold=True)
                items_run = para.add_run(", ".join(items))
                set_font(items_run, size=10.5)

    # --- EDUCATION ---
    education = data.get("education", [])
    if education:
        add_section_heading(doc, "Education")
        for edu in education:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(1)
            deg_run = para.add_run(edu.get("degree", ""))
            set_font(deg_run, size=11, bold=True)
            inst = edu.get("institution", "")
            dates = edu.get("dates", "")
            if inst:
                inst_run = para.add_run(f"  |  {inst}")
                set_font(inst_run, size=11)
            if dates:
                date_run = para.add_run(f"  |  {dates}")
                set_font(date_run, size=10, color=(89, 89, 89))
            notes = edu.get("notes", "").strip()
            if notes:
                note_para = doc.add_paragraph()
                note_para.paragraph_format.space_after = Pt(1)
                note_run = note_para.add_run(notes)
                set_font(note_run, size=10, color=(89, 89, 89))

    # --- CERTIFICATIONS ---
    certs = data.get("certifications", [])
    if certs:
        add_section_heading(doc, "Certifications")
        for cert in certs:
            add_bullet(doc, cert)

    doc.save(output_path)
    print(f"DOCX saved: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: build_docx.py <json_file> <output_docx>", file=sys.stderr)
        sys.exit(1)

    json_path = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.exists(json_path):
        print(f"Error: JSON file not found: {json_path}", file=sys.stderr)
        sys.exit(1)

    with open(json_path, "r", encoding="utf-8-sig") as f:
        data = json.load(f)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    build(data, output_path)


if __name__ == "__main__":
    main()
